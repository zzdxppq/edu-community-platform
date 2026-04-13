#!/usr/bin/env python3
"""Convert project-construction-plan.md to a styled Word document."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BRAND_COLOR = RGBColor(0x12, 0x83, 0xE9)
DARK_COLOR = RGBColor(0x0A, 0x5D, 0xAD)
HEADING_BG = "1283E9"
TABLE_HEADER_BG = "E8F4FD"
TABLE_ALT_BG = "F8FBFE"

def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for level in range(1, 5):
        h = doc.styles[f'Heading {level}']
        h.font.name = '微软雅黑'
        h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        h.font.color.rgb = DARK_COLOR if level <= 2 else BRAND_COLOR
        h.font.bold = True
        if level == 1:
            h.font.size = Pt(22)
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h.font.size = Pt(16)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        elif level == 3:
            h.font.size = Pt(13)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        else:
            h.font.size = Pt(11)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)


def set_cell_shading(cell, color):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_table(table, has_header=True):
    """Apply clean styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    for ci, cell in enumerate(table.rows[0].cells):
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x34, 0x40, 0x54)
        if has_header:
            set_cell_shading(cell, TABLE_HEADER_BG)

    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9.5)
        if ri > 0 and ri % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_ALT_BG)


def parse_table_row(line):
    """Parse a markdown table row into cells."""
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return cells


def is_separator_row(line):
    """Check if a line is a markdown table separator (e.g., |---|---|)."""
    return bool(re.match(r'^\|[\s\-:]+\|', line.strip()))


def add_formatted_text(paragraph, text):
    """Add text with inline bold/code formatting."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


def process_markdown(doc, md_text):
    """Parse markdown and build the Word document."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    skip_title = True  # skip the first H1 (document title handled separately)

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x36, 0x39, 0x3E)
                # Add light background via shading
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F6F8FA" w:val="clear"/>')
                p.paragraph_format.element.get_or_add_pPr().append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---' or stripped == '***':
            i += 1
            continue

        # Skip document end marker
        if stripped.startswith('*—') and stripped.endswith('—*'):
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)

            if level == 1 and skip_title:
                skip_title = False
                i += 1
                continue

            # Clean anchor links
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Table
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if not is_separator_row(row_line):
                    table_rows.append(parse_table_row(row_line))
                i += 1

            if len(table_rows) >= 1:
                num_cols = len(table_rows[0])
                # Ensure all rows have same column count
                for r in table_rows:
                    while len(r) < num_cols:
                        r.append('')
                    while len(r) > num_cols:
                        r.pop()

                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.autofit = True

                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        cell = table.cell(ri, ci)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        # Clean markdown formatting for display
                        clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                        clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)
                        add_formatted_text(p, cell_text)

                style_table(table)
                doc.add_paragraph()  # spacing after table
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x72, 0x85)
            run.font.size = Pt(10)
            i += 1
            continue

        # Unordered list
        list_match = re.match(r'^(\s*)[-*]\s+(.*)', stripped)
        if list_match:
            text = list_match.group(2)
            indent_level = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style='List Bullet')
            if indent_level > 2:
                p.paragraph_format.left_indent = Cm(1.5)
            p.clear()
            add_formatted_text(p, text)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1


def create_cover_page(doc):
    """Create a professional cover page."""
    # Add empty paragraphs for spacing
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('河南省县域城乡学校共同体发展平台')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK_COLOR
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('项 目 建 设 方 案')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BRAND_COLOR
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run('━' * 40)
    run.font.color.rgb = BRAND_COLOR
    run.font.size = Pt(12)

    # Version info
    for _ in range(4):
        doc.add_paragraph()

    info_items = [
        ('文档版本', 'V1.0'),
        ('编制日期', '2026年4月'),
        ('文档密级', '内部'),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x50, 0x55, 0x65)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()


def main():
    base = Path(__file__).parent
    md_path = base / 'project-construction-plan.md'
    out_path = base / '河南省县域城乡学校共同体发展平台_项目建设方案.docx'

    md_text = md_path.read_text(encoding='utf-8')

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    setup_styles(doc)
    create_cover_page(doc)
    process_markdown(doc, md_text)

    doc.save(str(out_path))
    print(f'Word document saved to: {out_path}')


if __name__ == '__main__':
    main()
